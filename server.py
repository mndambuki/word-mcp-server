from fastapi import FastAPI, HTTPException, Query, File, UploadFile, Request
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from typing import Union
from fastapi import Query, Body
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import io
import json
import uuid
from typing import Optional, Dict, Any, List
import uvicorn
import logging
import aiofiles
from markdownify import markdownify
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Word MCP Server", 
    version="1.0.0",
    description="A Word document generation MCP server"
)

# MCP Protocol support
MCP_VERSION = "2024-11-05"

# Create documents directory
DOCS_DIR = "/app/documents"
os.makedirs(DOCS_DIR, exist_ok=True)

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy", 
        "timestamp": datetime.utcnow().isoformat(),
        "service": "word-mcp-server",
        "port": 8004,
        "documents_directory": DOCS_DIR
    }

@app.post("/word/create")
async def create_word_document(
    # Query parameters (existing working method)
    title: str = Query(None, description="Document title"),
    content: str = Query(None, description="Document content"),
    author: str = Query("OpenWebUI User", description="Document author"),
    template: str = Query("standard", description="Document template"),
    # JSON body (for OpenWebUI compatibility)
    request_body: dict = Body(None)
):
    """Create a new Word document from text content"""
    try:
        # Use JSON body if provided, otherwise use query parameters
        if request_body:
            title = request_body.get("title", title)
            content = request_body.get("content", content) 
            author = request_body.get("author", author)
            template = request_body.get("template", template)
        
        if not title or not content:
            raise HTTPException(status_code=400, detail="Title and content are required")
            
        logger.info(f"Creating Word document: {title}")
        
        # Create a new document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        
        # Apply template styling
        apply_template_styling(doc, template)
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Created: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Author: {author}")
        doc.add_paragraph("")  # Empty line
        
        # Process and add content
        process_content_to_document(doc, content)
        
        # Generate unique filename
        doc_id = str(uuid.uuid4())[:8]
        safe_title = re.sub(r'[^\w\s-]', '', title).strip()[:50]
        filename = f"{safe_title}_{doc_id}.docx"
        filepath = os.path.join(DOCS_DIR, filename)
        
        # Save document
        doc.save(filepath)
        
        result = {
            "success": True,
            "document_id": doc_id,
            "filename": filename,
            "title": title,
            "author": author,
            "template": template,
            "file_path": filepath,
            "file_size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat(),
            "download_url": f"/word/download/{filename}"
        }
        
        logger.info(f"Successfully created document: {filename}")
        return result
        
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating document: {str(e)}")

@app.post("/word/create-from-chat")
async def create_document_from_chat(
    chat_title: str = Query(..., description="Chat conversation title"),
    messages: str = Query(..., description="JSON string of chat messages"),
    format_style: str = Query("conversation", description="Format: conversation, summary, report")
):
    """Create a Word document from chat conversation"""
    try:
        logger.info(f"Creating document from chat: {chat_title}")
        
        # Parse messages
        try:
            chat_messages = json.loads(messages)
        except json.JSONDecodeError:
            # If not JSON, treat as plain text
            chat_messages = [{"role": "user", "content": messages}]
        
        # Create document
        doc = Document()
        
        # Set properties
        doc.core_properties.title = f"Chat: {chat_title}"
        doc.core_properties.author = "OpenWebUI"
        doc.core_properties.created = datetime.now()
        
        # Add title
        title_paragraph = doc.add_heading(f"Chat Conversation: {chat_title}", 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Messages: {len(chat_messages) if isinstance(chat_messages, list) else 1}")
        doc.add_paragraph("")
        
        # Format based on style
        if format_style == "conversation":
            format_as_conversation(doc, chat_messages)
        elif format_style == "summary":
            format_as_summary(doc, chat_messages)
        else:  # report
            format_as_report(doc, chat_messages)
        
        # Generate filename
        doc_id = str(uuid.uuid4())[:8]
        safe_title = re.sub(r'[^\w\s-]', '', chat_title).strip()[:50]
        filename = f"Chat_{safe_title}_{doc_id}.docx"
        filepath = os.path.join(DOCS_DIR, filename)
        
        # Save document
        doc.save(filepath)
        
        result = {
            "success": True,
            "document_id": doc_id,
            "filename": filename,
            "title": f"Chat: {chat_title}",
            "format_style": format_style,
            "message_count": len(chat_messages) if isinstance(chat_messages, list) else 1,
            "file_path": filepath,
            "file_size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat(),
            "download_url": f"/word/download/{filename}"
        }
        
        logger.info(f"Successfully created chat document: {filename}")
        return result
        
    except Exception as e:
        logger.error(f"Error creating chat document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error creating chat document: {str(e)}")

@app.get("/word/download/{filename}")
async def download_document(filename: str):
    """Download a Word document"""
    try:
        filepath = os.path.join(DOCS_DIR, filename)
        
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="Document not found")
        
        return FileResponse(
            filepath,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error downloading document: {str(e)}")

@app.get("/word/list")
async def list_documents():
    """List all created documents"""
    try:
        documents = []
        
        for filename in os.listdir(DOCS_DIR):
            if filename.endswith('.docx'):
                filepath = os.path.join(DOCS_DIR, filename)
                stat = os.stat(filepath)
                
                documents.append({
                    "filename": filename,
                    "size": stat.st_size,
                    "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "download_url": f"/word/download/{filename}"
                })
        
        # Sort by creation time, newest first
        documents.sort(key=lambda x: x['created'], reverse=True)
        
        return {
            "success": True,
            "document_count": len(documents),
            "documents": documents
        }
        
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error listing documents: {str(e)}")

@app.delete("/word/delete/{filename}")
async def delete_document(filename: str):
    """Delete a Word document"""
    try:
        filepath = os.path.join(DOCS_DIR, filename)
        
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="Document not found")
        
        os.remove(filepath)
        
        return {
            "success": True,
            "message": f"Document {filename} deleted successfully"
        }
        
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

def apply_template_styling(doc, template):
    """Apply styling based on template"""
    styles = doc.styles
    
    if template == "report":
        # Professional report styling
        try:
            heading_style = styles['Heading 1']
            heading_style.font.size = Pt(16)
            heading_style.font.bold = True
        except:
            pass
    elif template == "memo":
        # Memo styling
        try:
            normal_style = styles['Normal']
            normal_style.font.size = Pt(11)
        except:
            pass
    # Standard template is default

def process_content_to_document(doc, content):
    """Process content and add to document with formatting"""
    
    # Split content into lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph("")
            continue
            
        # Check for markdown-style headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            # Numbered list
            p = doc.add_paragraph(line[3:], style='List Number')
        else:
            # Regular paragraph
            doc.add_paragraph(line)

def format_as_conversation(doc, messages):
    """Format messages as a conversation"""
    doc.add_heading("Conversation", level=1)
    
    for i, msg in enumerate(messages):
        if isinstance(msg, dict):
            role = msg.get('role', 'unknown').title()
            content = msg.get('content', '')
        else:
            role = "Message"
            content = str(msg)
        
        # Add role header
        role_paragraph = doc.add_paragraph()
        role_run = role_paragraph.add_run(f"{role}:")
        role_run.bold = True
        role_run.font.size = Pt(12)
        
        # Add content
        doc.add_paragraph(content)
        doc.add_paragraph("")  # Empty line between messages

def format_as_summary(doc, messages):
    """Format messages as a summary"""
    doc.add_heading("Summary", level=1)
    
    # Extract key information
    user_messages = []
    assistant_messages = []
    
    for msg in messages:
        if isinstance(msg, dict):
            if msg.get('role') == 'user':
                user_messages.append(msg.get('content', ''))
            elif msg.get('role') == 'assistant':
                assistant_messages.append(msg.get('content', ''))
    
    if user_messages:
        doc.add_heading("User Queries", level=2)
        for i, query in enumerate(user_messages, 1):
            doc.add_paragraph(f"{i}. {query}")
        doc.add_paragraph("")
    
    if assistant_messages:
        doc.add_heading("Key Responses", level=2)
        for i, response in enumerate(assistant_messages, 1):
            doc.add_paragraph(f"{i}. {response[:200]}..." if len(response) > 200 else f"{i}. {response}")

def format_as_report(doc, messages):
    """Format messages as a structured report"""
    doc.add_heading("Report", level=1)
    
    doc.add_heading("Overview", level=2)
    doc.add_paragraph(f"This report contains {len(messages)} messages from a chat conversation.")
    
    doc.add_heading("Detailed Content", level=2)
    
    for i, msg in enumerate(messages, 1):
        if isinstance(msg, dict):
            role = msg.get('role', 'unknown').title()
            content = msg.get('content', '')
        else:
            role = "Content"
            content = str(msg)
        
        doc.add_heading(f"Section {i}: {role}", level=3)
        doc.add_paragraph(content)

@app.get("/")
async def root():
    """Root endpoint with service information"""
    return {
        "service": "Word MCP Server",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "create_document": "/word/create",
            "create_from_chat": "/word/create-from-chat",
            "download": "/word/download/{filename}",
            "list_documents": "/word/list",
            "delete_document": "/word/delete/{filename}",
            "mcp": "/mcp"
        },
        "supported_formats": ["docx"],
        "templates": ["standard", "report", "memo", "letter"],
        "mcp_version": MCP_VERSION,
        "timestamp": datetime.utcnow().isoformat()
    }

# MCP Protocol endpoint (JSON-RPC 2.0)
@app.post("/mcp")
async def handle_mcp_protocol(request: Request):
    """Handle MCP protocol requests (JSON-RPC 2.0)"""
    try:
        body = await request.json()
        
        # Handle both single requests and batch requests
        if isinstance(body, list):
            responses = [process_mcp_request(req) for req in body]
            return responses
        else:
            response = process_mcp_request(body)
            if response is None:
                # Notification request (no response needed)
                return JSONResponse(status_code=204, content={})
            return response
    
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={
                "jsonrpc": "2.0",
                "error": {
                    "code": -32700,
                    "message": "Parse error"
                },
                "id": None
            }
        )
    except Exception as e:
        logger.error(f"MCP Protocol error: {str(e)}")
        return JSONResponse(
            status_code=500,
            content={
                "jsonrpc": "2.0",
                "error": {
                    "code": -32603,
                    "message": f"Internal error: {str(e)}"
                },
                "id": None
            }
        )

def process_mcp_request(request_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """Process a single MCP JSON-RPC request"""
    
    # Validate JSON-RPC structure
    if not isinstance(request_data, dict) or request_data.get("jsonrpc") != "2.0":
        return {
            "jsonrpc": "2.0",
            "error": {
                "code": -32600,
                "message": "Invalid Request"
            },
            "id": request_data.get("id")
        }
    
    method = request_data.get("method")
    params = request_data.get("params", {})
    request_id = request_data.get("id")
    
    logger.info(f"MCP method: {method}, params: {params}")
    
    # Handle different MCP methods
    if method == "initialize":
        result = handle_initialize(params)
    elif method == "tools/list":
        result = handle_tools_list()
    elif method == "tools/call":
        result = handle_tools_call(params)
    elif method == "resources/list":
        result = handle_resources_list()
    elif method == "resources/read":
        result = handle_resources_read(params)
    else:
        return {
            "jsonrpc": "2.0",
            "error": {
                "code": -32601,
                "message": f"Method not found: {method}"
            },
            "id": request_id
        }
    
    # Notification requests have no id and should not return a response
    if request_id is None:
        return None
    
    if isinstance(result, dict) and "error" in result:
        return {
            "jsonrpc": "2.0",
            "error": result["error"],
            "id": request_id
        }
    
    return {
        "jsonrpc": "2.0",
        "result": result,
        "id": request_id
    }

def handle_initialize(params: Dict[str, Any]) -> Dict[str, Any]:
    """Handle MCP initialize request"""
    return {
        "protocolVersion": MCP_VERSION,
        "capabilities": {
            "tools": {}
        },
        "serverInfo": {
            "name": "Word MCP Server",
            "version": "1.0.0"
        }
    }

def handle_tools_list() -> Dict[str, Any]:
    """Return list of available tools"""
    return {
        "tools": [
            {
                "name": "create_document",
                "description": "Create a Word document from text content with markdown support",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "Document title"
                        },
                        "content": {
                            "type": "string",
                            "description": "Document content (supports markdown)"
                        },
                        "author": {
                            "type": "string",
                            "description": "Document author",
                            "default": "Word MCP Server"
                        },
                        "template": {
                            "type": "string",
                            "enum": ["standard", "report", "memo", "letter"],
                            "description": "Document template",
                            "default": "standard"
                        }
                    },
                    "required": ["title", "content"]
                }
            },
            {
                "name": "create_from_chat",
                "description": "Create a Word document from a chat conversation",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "chat_title": {
                            "type": "string",
                            "description": "Chat conversation title"
                        },
                        "messages": {
                            "type": "array",
                            "description": "Array of chat messages",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "role": {
                                        "type": "string",
                                        "enum": ["user", "assistant", "system"]
                                    },
                                    "content": {
                                        "type": "string"
                                    }
                                }
                            }
                        },
                        "format_style": {
                            "type": "string",
                            "enum": ["conversation", "summary", "report"],
                            "description": "Format style for the document",
                            "default": "conversation"
                        }
                    },
                    "required": ["chat_title", "messages"]
                }
            },
            {
                "name": "list_documents",
                "description": "List all created documents",
                "inputSchema": {
                    "type": "object",
                    "properties": {}
                }
            },
            {
                "name": "delete_document",
                "description": "Delete a Word document by filename",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "filename": {
                            "type": "string",
                            "description": "Filename to delete"
                        }
                    },
                    "required": ["filename"]
                }
            }
        ]
    }

def handle_tools_call(params: Dict[str, Any]) -> Dict[str, Any]:
    """Handle tool call requests"""
    tool_name = params.get("name")
    tool_params = params.get("arguments", {})
    
    logger.info(f"Tool call: {tool_name} with params: {tool_params}")
    
    try:
        if tool_name == "create_document":
            return call_create_document(tool_params)
        elif tool_name == "create_from_chat":
            return call_create_from_chat(tool_params)
        elif tool_name == "list_documents":
            return call_list_documents()
        elif tool_name == "delete_document":
            return call_delete_document(tool_params)
        else:
            return {
                "error": {
                    "code": -32601,
                    "message": f"Tool not found: {tool_name}"
                }
            }
    except Exception as e:
        logger.error(f"Error calling tool {tool_name}: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error calling tool: {str(e)}"
            }
        }

def call_create_document(params: Dict[str, Any]) -> Dict[str, Any]:
    """Create a Word document (tool call wrapper)"""
    title = params.get("title")
    content = params.get("content")
    author = params.get("author", "Word MCP Server")
    template = params.get("template", "standard")
    
    if not title or not content:
        return {
            "error": {
                "code": -32602,
                "message": "Title and content are required"
            }
        }
    
    try:
        # Create a new document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()
        
        # Apply template styling
        apply_template_styling(doc, template)
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Created: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Author: {author}")
        doc.add_paragraph("")  # Empty line
        
        # Process and add content
        process_content_to_document(doc, content)
        
        # Generate unique filename
        doc_id = str(uuid.uuid4())[:8]
        safe_title = re.sub(r'[^\w\s-]', '', title).strip()[:50]
        filename = f"{safe_title}_{doc_id}.docx"
        filepath = os.path.join(DOCS_DIR, filename)
        
        # Save document
        doc.save(filepath)
        
        return {
            "document_id": doc_id,
            "filename": filename,
            "title": title,
            "author": author,
            "template": template,
            "file_path": filepath,
            "file_size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error creating document: {str(e)}"
            }
        }

def call_create_from_chat(params: Dict[str, Any]) -> Dict[str, Any]:
    """Create a document from chat (tool call wrapper)"""
    chat_title = params.get("chat_title")
    messages = params.get("messages", [])
    format_style = params.get("format_style", "conversation")
    
    if not chat_title or not messages:
        return {
            "error": {
                "code": -32602,
                "message": "chat_title and messages are required"
            }
        }
    
    try:
        doc = Document()
        doc.core_properties.title = f"Chat: {chat_title}"
        doc.core_properties.author = "Word MCP Server"
        doc.core_properties.created = datetime.now()
        
        # Add title
        title_paragraph = doc.add_heading(f"Chat Conversation: {chat_title}", 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Messages: {len(messages)}")
        doc.add_paragraph("")
        
        # Format based on style
        if format_style == "conversation":
            format_as_conversation(doc, messages)
        elif format_style == "summary":
            format_as_summary(doc, messages)
        else:  # report
            format_as_report(doc, messages)
        
        # Generate filename
        doc_id = str(uuid.uuid4())[:8]
        safe_title = re.sub(r'[^\w\s-]', '', chat_title).strip()[:50]
        filename = f"Chat_{safe_title}_{doc_id}.docx"
        filepath = os.path.join(DOCS_DIR, filename)
        
        # Save document
        doc.save(filepath)
        
        return {
            "document_id": doc_id,
            "filename": filename,
            "title": f"Chat: {chat_title}",
            "format_style": format_style,
            "message_count": len(messages),
            "file_path": filepath,
            "file_size": os.path.getsize(filepath),
            "created_at": datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error creating chat document: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error creating chat document: {str(e)}"
            }
        }

def call_list_documents() -> Dict[str, Any]:
    """List documents (tool call wrapper)"""
    try:
        documents = []
        
        for filename in os.listdir(DOCS_DIR):
            if filename.endswith('.docx'):
                filepath = os.path.join(DOCS_DIR, filename)
                stat = os.stat(filepath)
                
                documents.append({
                    "filename": filename,
                    "size": stat.st_size,
                    "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
        
        documents.sort(key=lambda x: x['created'], reverse=True)
        
        return {
            "document_count": len(documents),
            "documents": documents
        }
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error listing documents: {str(e)}"
            }
        }

def call_delete_document(params: Dict[str, Any]) -> Dict[str, Any]:
    """Delete a document (tool call wrapper)"""
    filename = params.get("filename")
    
    if not filename:
        return {
            "error": {
                "code": -32602,
                "message": "filename is required"
            }
        }
    
    try:
        filepath = os.path.join(DOCS_DIR, filename)
        
        if not os.path.exists(filepath):
            return {
                "error": {
                    "code": -32603,
                    "message": f"Document not found: {filename}"
                }
            }
        
        os.remove(filepath)
        
        return {
            "message": f"Document {filename} deleted successfully"
        }
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error deleting document: {str(e)}"
            }
        }

def handle_resources_list() -> Dict[str, Any]:
    """Handle resources list request"""
    try:
        documents = []
        
        for filename in os.listdir(DOCS_DIR):
            if filename.endswith('.docx'):
                filepath = os.path.join(DOCS_DIR, filename)
                
                documents.append({
                    "uri": f"word://{filename}",
                    "name": filename,
                    "description": f"Word document: {filename}",
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                })
        
        return {
            "resources": documents
        }
    except Exception as e:
        logger.error(f"Error listing resources: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error listing resources: {str(e)}"
            }
        }

def handle_resources_read(params: Dict[str, Any]) -> Dict[str, Any]:
    """Handle resource read request"""
    uri = params.get("uri")
    
    if not uri or not uri.startswith("word://"):
        return {
            "error": {
                "code": -32602,
                "message": "Invalid resource URI"
            }
        }
    
    filename = uri.replace("word://", "")
    
    try:
        filepath = os.path.join(DOCS_DIR, filename)
        
        if not os.path.exists(filepath):
            return {
                "error": {
                    "code": -32603,
                    "message": f"Resource not found: {uri}"
                }
            }
        
        with open(filepath, 'rb') as f:
            content = f.read()
        
        # Return base64 encoded content or file info
        import base64
        return {
            "contents": [
                {
                    "uri": uri,
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "text": f"Word document: {filename} ({len(content)} bytes)"
                }
            ]
        }
    except Exception as e:
        logger.error(f"Error reading resource: {str(e)}")
        return {
            "error": {
                "code": -32603,
                "message": f"Error reading resource: {str(e)}"
            }
        }

@app.get("/tools")
async def get_tools():
    """Get available tools for MCP proxy"""
    return {
        "tools": [
            {
                "name": "create_document",
                "description": "Create a Word document from text content",
                "endpoint": "/word/create",
                "method": "POST",
                "parameters": {
                    "title": {
                        "type": "string",
                        "description": "Document title",
                        "required": True
                    },
                    "content": {
                        "type": "string", 
                        "description": "Document content (supports markdown)",
                        "required": True
                    },
                    "author": {
                        "type": "string",
                        "description": "Document author",
                        "default": "OpenWebUI User",
                        "required": False
                    },
                    "template": {
                        "type": "string",
                        "description": "Document template: standard, report, memo, letter", 
                        "default": "standard",
                        "required": False
                    }
                }
            },
            {
                "name": "create_from_chat",
                "description": "Create a Word document from chat conversation",
                "endpoint": "/word/create-from-chat", 
                "method": "POST",
                "parameters": {
                    "chat_title": {
                        "type": "string",
                        "description": "Chat conversation title",
                        "required": True
                    },
                    "messages": {
                        "type": "string",
                        "description": "Chat messages as JSON string",
                        "required": True
                    },
                    "format_style": {
                        "type": "string",
                        "description": "Format style: conversation, summary, report",
                        "default": "conversation",
                        "required": False
                    }
                }
            }
        ]
    }

if __name__ == "__main__":
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8004,
        log_level="info"
    )
